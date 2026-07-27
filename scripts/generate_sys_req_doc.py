import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fi ll_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def build_sys_req_document(output_path):
    doc = Document()
    
    # Page Setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    # Colors
    NAVY_HEX = "1B365D"
    NAVY_RGB = RGBColor(0x1B, 0x36, 0x5D)
    TEAL_HEX = "008080"
    TEAL_RGB = RGBColor(0x00, 0x80, 0x80)
    GREY_HEX = "555555"
    LIGHT_BG = "F4F6F9"
    
    # --- Title Banner ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("FinAuditPro")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY_RGB
    
    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run("Real-World Production System Requirements Specification")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = TEAL_RGB
    
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run("Document Version: 1.0.0  |  Classification: Executive / Technical  |  Date: July 2026")
    meta_run.font.name = 'Calibri'
    meta_run.font.size = Pt(9.5)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Helper functions for structured headings & text
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY_RGB
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = TEAL_RGB
        return h

    def add_callout(text, bold_prefix="NOTE: "):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, LIGHT_BG)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.color.rgb = NAVY_RGB
        r2 = p.add_run(text)
        r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- Section 1: Executive Overview ---
    add_heading_1("1. Executive Overview")
    p = doc.add_paragraph(
        "FinAuditPro is a next-generation AI-powered executive intelligence platform designed specifically for Chartered Accountants (CAs), "
        "statutory auditors, and enterprise financial compliance teams. Built with a 100% air-gapped, offline-first philosophy, FinAuditPro guarantees "
        "that sensitive financial ledgers, tax filings, and client audit data never transmit across public cloud networks."
    )
    p.paragraph_format.space_after = Pt(8)
    
    p2 = doc.add_paragraph(
        "This document defines the real-world infrastructure, hardware, software, security, network, and regulatory compliance specifications "
        "necessary to successfully deploy FinAuditPro across desktop single-user workstations and multi-auditor enterprise server environments."
    )
    p2.paragraph_format.space_after = Pt(12)

    add_callout(
        "FinAuditPro executes local Retrieval-Augmented Generation (RAG) using Ollama (Llama 3.2 / DeepSeek-R1), local FAISS vector search, "
        "and multi-engine OCR entirely on on-premise hardware without external API key dependencies.",
        "KEY ARCHITECTURAL HIGHLIGHT: "
    )

    # --- Section 2: Deployment Topologies ---
    add_heading_1("2. Deployment Topologies")
    
    add_heading_2("2.1 Option A: Standalone Air-Gapped Desktop Workstation")
    p = doc.add_paragraph(
        "Ideal for independent Chartered Accountants, sole practitioners, and small auditing firms operating on standalone laptops or workstations. "
        "All components—including the PySide6 Qt GUI, SQLite WAL database, local vector store (FAISS), and Ollama AI daemon—run on a single machine."
    )
    p.paragraph_format.space_after = Pt(8)

    add_heading_2("2.2 Option B: Enterprise Multi-User Intranet Server")
    p = doc.add_paragraph(
        "Engineered for corporate audit departments and large CA firms with multiple concurrent auditors. Client desktop instances connect over a secure local area network (LAN/VLAN) "
        "to a central PostgreSQL database server and a dedicated GPU server hosting the Ollama local AI daemon and shared FAISS vector indices."
    )
    p.paragraph_format.space_after = Pt(12)

    # --- Section 3: Hardware Specifications ---
    add_heading_1("3. Hardware Specifications")

    add_heading_2("3.1 Client Workstation Hardware Requirements (Auditor Laptops/Desktops)")
    
    # Table 1: Client Hardware
    t1 = doc.add_table(rows=6, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)
    
    headers1 = ["Hardware Component", "Minimum Requirement", "Recommended (Optimal Performance)"]
    for i, h in enumerate(headers1):
        cell = t1.cell(0, i)
        set_cell_background(cell, NAVY_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data1 = [
        ("Processor (CPU)", "Intel Core i5 (10th Gen+) / AMD Ryzen 5 3000+ (4 cores, 2.5 GHz+)", "Intel Core i7/i9 (12th Gen+) / AMD Ryzen 7/9 / Apple Silicon (M1/M2/M3/M4 Pro/Max)"),
        ("System Memory (RAM)", "16 GB DDR4 (Mandatory for local LLM inference)", "32 GB – 64 GB DDR4/DDR5 / Apple Unified Memory"),
        ("Storage (Disk)", "20 GB free space on SSD (SATA or NVMe)", "100 GB+ NVMe M.2 SSD (Read speed ≥ 3500 MB/s)"),
        ("Graphics (GPU)", "Integrated Intel Iris Xe / AMD Radeon", "Dedicated NVIDIA GPU with 8GB+ VRAM (RTX 3060/4060+) or Apple Silicon GPU"),
        ("Display Resolution", "1920 × 1080 (Full HD)", "Dual Monitors / 2K+ (2560 × 1440) resolution")
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        bg = LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            cell = t1.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            r = p.add_run(cell_value)
            if col_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    add_heading_2("3.2 Central AI & Analytics Server Hardware (Enterprise Tier)")

    # Table 2: Server Hardware
    t2 = doc.add_table(rows=6, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)

    headers2 = ["Server Component", "Minimum Server Spec (7B–8B LLM Models)", "Recommended Server Spec (14B–70B Models & Multi-User)"]
    for i, h in enumerate(headers2):
        cell = t2.cell(0, i)
        set_cell_background(cell, NAVY_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data2 = [
        ("Server CPU", "Intel Xeon Silver / AMD EPYC (8+ Cores)", "Intel Xeon Gold / AMD EPYC (16–32 Cores)"),
        ("System RAM", "32 GB DDR4 ECC RAM", "128 GB+ DDR5 ECC RAM"),
        ("GPU Acceleration", "1× NVIDIA RTX 4090 / A4000 (16GB–24GB VRAM)", "2×–4× NVIDIA A100 / H100 / RTX 6000 Ada (48GB+ VRAM total)"),
        ("Storage Subsystem", "500 GB Enterprise NVMe SSD (RAID-1)", "2 TB RAID-10 NVMe Enterprise SSD"),
        ("Network Controller", "1 Gbps Enterprise NIC", "10 Gbps Enterprise Dual-Port NIC")
    ]

    for row_idx, row_data in enumerate(data2, start=1):
        bg = LIGHT_BG if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, cell_value in enumerate(row_data):
            cell = t2.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            p = cell.paragraphs[0]
            r = p.add_run(cell_value)
            if col_idx == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Section 4: Software Stack ---
    add_heading_1("4. Software Stack & Environment Requirements")
    
    bullet_items = [
        ("Operating Systems: ", "Windows 10/11 (64-bit, Build 19041+), Windows Server 2022, macOS 12+ (Apple Silicon/Intel), or Linux (Ubuntu 22.04/24.04 LTS, RHEL 8/9)."),
        ("Python Core Runtime: ", "Python 3.11 or 3.12 (64-bit)."),
        ("GUI Framework: ", "PySide6 (Qt 6.7 / 6.8+)."),
        ("Database Systems: ", "SQLite 3.35+ with Write-Ahead Logging (WAL) for Desktop; PostgreSQL 15/16 for Enterprise deployment."),
        ("AI Inference Engine: ", "Ollama Daemon v0.3.3+ executing models such as llama3.2 (3B/8B), deepseek-r1 (7B/8B/14B), or mistral-7b."),
        ("Vector Search & Embeddings: ", "FAISS (faiss-cpu / faiss-gpu v1.8.0) and SentenceTransformers (all-MiniLM-L6-v2 / bge-small-en-v1.5)."),
        ("Document Parsing & OCR: ", "PyPDF (v4.3+), pdfplumber (v0.11+), Tesseract OCR (v5.3+), and PaddleOCR (v2.7+)."),
        ("Reporting & Exporters: ", "ReportLab (v4.2+) for digitally signed PDF Audit Packs; OpenPyXL (v3.1+) for Excel financial ledgers.")
    ]

    for bold_prefix, text in bullet_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.color.rgb = NAVY_RGB
        r2 = bp.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- Section 5: Security & Network Architecture ---
    add_heading_1("5. Security, Network & Infrastructure Architecture")
    
    p = doc.add_paragraph(
        "FinAuditPro incorporates end-to-end security safeguards to protect client financial integrity and prevent data exposure:"
    )
    p.paragraph_format.space_after = Pt(6)

    sec_items = [
        ("Air-Gapped Isolation: ", "Operates with zero outbound internet network calls. All processing is localized within the company intranet or desktop sandbox."),
        ("Port Allocation & Binding: ", "11434/TCP (Ollama local REST API, bound to 127.0.0.1 by default); 5432/TCP (PostgreSQL central connection)."),
        ("Data Encryption at Rest: ", "AES-256-GCM encryption for database backup archives and export payloads."),
        ("Credential Protection: ", "PBKDF2-HMAC-SHA256 password hashing configured with 100,000+ iterations."),
        ("Audit Log Hash Chain: ", "Immutable cryptographic ledger utilizing SHA-256 hash-chaining (Hash_n = SHA256(Hash_{n-1} + Entry_Data)) to prevent audit log tampering.")
    ]

    for bold_prefix, text in sec_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.color.rgb = NAVY_RGB
        r2 = bp.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- Section 6: Compliance & Regulatory Readiness ---
    add_heading_1("6. Compliance & Regulatory Readiness")
    
    comp_items = [
        ("Statutory Financial Rules: ", "Deterministic rule engines for GSTIN checksums, PAN syntax, Income Tax Section 40A(3) cash expenditure limits, and Benford's Law anomaly detection."),
        ("Data Protection Laws: ", "Fully compliant with EU GDPR, Indian DPDP Act 2023, and US HIPAA financial data guidelines due to zero third-party cloud data transmission."),
        ("ISO 27001 Readiness: ", "Meets ISO 27001 Access Control (A.9) and Cryptography (A.10) standard controls.")
    ]

    for bold_prefix, text in comp_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.color.rgb = NAVY_RGB
        r2 = bp.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Section 7: Packaging & Deployment ---
    add_heading_1("7. Distribution, Packaging & Operational Setup")
    
    p = doc.add_paragraph(
        "FinAuditPro provides simplified zero-configuration installation bootstrappers and standalone compilations:"
    )
    p.paragraph_format.space_after = Pt(6)

    pkg_items = [
        ("Windows Executable: ", "Compiled into a standalone directory using PyInstaller (FinAuditPro.spec) and wrapped into a single setup installer (FinAuditPro_v1.0.0_Setup.exe) using Inno Setup 6."),
        ("Auto-Installer Scripts: ", "Includes install.bat (Windows) and install.sh (macOS/Linux) which detect runtime environments, build virtualenv, install dependencies, and launch."),
        ("Ollama Initialization: ", "Pre-flight command required prior to first execution: ollama pull llama3.2")
    ]

    for bold_prefix, text in pkg_items:
        bp = doc.add_paragraph(style='List Bullet')
        bp.paragraph_format.space_after = Pt(4)
        r1 = bp.add_run(bold_prefix)
        r1.font.bold = True
        r1.font.color.rgb = NAVY_RGB
        r2 = bp.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # --- Footer note ---
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer_p.add_run("— End of Production System Requirements Specification —")
    frun.font.italic = True
    frun.font.size = Pt(10)
    frun.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Save document
    doc.save(output_path)
    print(f"Document saved successfully to {output_path}")

if __name__ == "__main__":
    build_sys_req_document("FinAuditPro_System_Requirements.docx")
