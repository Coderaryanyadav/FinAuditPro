"""
Unified Document Parser Module for FinAuditPro.
Provides format readers for PDF, Excel, CSV, Images, Word (.docx), TXT, XML, and JSON files.
"""

from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
import os
import csv
import logging

from .ocr_engine import OCREngine, OCRResult
from .table_extractor import TableExtractor, ExtractedTable
from .text_cleaner import TextCleaner

logger = logging.getLogger(__name__)

@dataclass
class ParsedDocument:
    file_path: str
    file_name: str
    file_type: str
    file_size_bytes: int
    raw_text: str
    cleaned_text: str
    pages_count: int
    tables: List[ExtractedTable] = field(default_factory=list)
    ocr_result: Optional[OCRResult] = None


class DocumentParser:
    """Unified document parser for multi-format financial files."""

    @classmethod
    def parse_document(cls, file_path: str) -> ParsedDocument:
        """Parses any supported file format into a unified ParsedDocument representation."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Cannot parse non-existent file: {file_path}")

        file_name = os.path.basename(file_path)
        file_size = os.path.getsize(file_path)
        ext = os.path.splitext(file_path)[1].lower()

        raw_text = ""
        tables: List[ExtractedTable] = []
        pages_count = 1
        ocr_res = None

        if ext == ".pdf":
            ocr_res = OCREngine.process_document(file_path)
            raw_text = ocr_res.raw_text
            pages_count = ocr_res.total_pages
            tables = TableExtractor.extract_tables_from_pdf(file_path)

        elif ext in [".xlsx", ".xls"]:
            tables = TableExtractor.extract_tables_from_excel(file_path)
            text_parts = []
            for tbl in tables:
                text_parts.append(" | ".join(tbl.headers))
                for r in tbl.rows:
                    text_parts.append(" | ".join(r))
            raw_text = "\n".join(text_parts)
            pages_count = len(tables)

        elif ext == ".csv":
            tables = TableExtractor.extract_tables_from_csv(file_path)
            if tables:
                headers = " | ".join(tables[0].headers)
                rows = [" | ".join(r) for r in tables[0].rows]
                raw_text = f"{headers}\n" + "\n".join(rows)

        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
            ocr_res = OCREngine.process_ocr_fallback(file_path)
            raw_text = ocr_res.raw_text

        elif ext == ".docx":
            try:
                import docx
                doc = docx.Document(file_path)
                raw_text = "\n".join([p.text for p in doc.paragraphs if p.text])
            except (OSError, ValueError, RuntimeError) as e:
                logger.error(f"Docx parsing failed for {file_path}: {e}")
                raw_text = ""

        elif ext in [".txt", ".json", ".xml"]:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                raw_text = f.read()

        cleaned = TextCleaner.clean_text(raw_text)

        return ParsedDocument(
            file_path=file_path,
            file_name=file_name,
            file_type=ext.replace(".", "").upper(),
            file_size_bytes=file_size,
            raw_text=raw_text,
            cleaned_text=cleaned,
            pages_count=pages_count,
            tables=tables,
            ocr_result=ocr_res
        )
